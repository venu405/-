"""文档解析与分块（入库链路第一步）。

支持格式（P1 起步）：Markdown / 纯文本 / PDF / Word(docx)
后续可扩展：HTML、Excel、图片 OCR

分块策略（P1 简单版）：按字符固定大小 + 重叠窗口，保留元数据。
注意：分块质量直接影响检索质量——这是 RAG 的"GIGO"关卡。
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


@dataclass(kw_only=True)
class DocumentChunk:
    """一个分块：检索的最小单元，携带定位元数据。"""

    text: str
    doc_id: str            # 所属文档 ID（UUID）
    doc_title: str         # 文档标题
    chunk_index: int       # 第几个分块（0 起）
    source_type: str       # pdf / markdown / text / docx
    kb_id: str = "default"   # 所属知识库（多库隔离，P3）
    page: int | None = None   # 页码（PDF 有，其余 None）
    metadata: dict[str, Any] = field(default_factory=dict)


def parse_document(path: Path) -> str:
    """按扩展名解析文档 → 纯文本。解析失败抛异常（由上层降级处理）。"""
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型: {ext}，支持 {sorted(SUPPORTED_EXTENSIONS)}")

    if ext == ".md" or ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext == ".docx":
        return _parse_docx(path)
    raise ValueError(f"未实现的解析器: {ext}")


def _parse_pdf(path: Path) -> str:
    """PDF 解析：逐页提取文本，页间用换行分隔。"""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _parse_docx(path: Path) -> str:
    """Word 解析：提取段落 + 表格。"""
    import docx

    doc = docx.Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def chunk_text(
    text: str,
    *,
    chunk_size: int = 800,
    overlap: int = 100,
) -> list[str]:
    """按字符分块：固定大小 + 重叠窗口。

    overlap 的作用：防止"语义在边界被切断"（比如一句话正好被切成两半）。
    代价：内容重复存储，但检索召回率更高——这是 RAG 的标准取舍。
    """
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []

    chunks: list[str] = []
    step = max(chunk_size - overlap, 1)  # 每次前进步长
    start = 0
    while start < len(text):
        chunks.append(text[start : start + chunk_size])
        start += step
        # 防死循环：最后一段不足时退出
        if start >= len(text):
            break
    return chunks


def build_chunks(
    path: Path,
    *,
    doc_id: str,
    chunk_size: int = 800,
    overlap: int = 100,
    kb_id: str = "default",
) -> list[DocumentChunk]:
    """完整分块管线：解析 → 分块 → 挂元数据。kb_id 标记所属知识库（P3）。"""
    text = parse_document(path)
    parts = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    return [
        DocumentChunk(
            text=part,
            doc_id=doc_id,
            doc_title=path.stem,
            chunk_index=i,
            source_type=path.suffix.lstrip("."),
            kb_id=kb_id,
        )
        for i, part in enumerate(parts)
    ]
